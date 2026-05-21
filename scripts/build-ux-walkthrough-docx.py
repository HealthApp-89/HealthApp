"""
Build the UX walkthrough 1-pager as a Word doc.
Usage: python3 scripts/build-ux-walkthrough-docx.py
Output: docs/Apex-Health-OS-UX-Walkthrough.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parents[1] / "docs" / "Apex-Health-OS-UX-Walkthrough.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# Tighten margins so it fits on one page
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    p.paragraph_format.space_after = Pt(4)


def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def add_bullet(parts, level=0):
    """parts: list of (text, {'bold': bool, 'italic': bool}) tuples, or a plain string."""
    style_name = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if isinstance(parts, str):
        parts = [(parts, {})]
    for text, opts in parts:
        run = p.add_run(text)
        run.font.size = Pt(10)
        if opts.get("bold"):
            run.bold = True
        if opts.get("italic"):
            run.italic = True


def b(text):
    return (text, {"bold": True})


def t(text):
    return (text, {})


def i(text):
    return (text, {"italic": True})


# --- Title
add_title("Apex Health OS — End-to-End User Experience")

# --- 1. First open
add_h2("1. First open (one-time, ~15–20 min)")
add_bullet([b("Sign up / log in"), t(" on /login (Supabase Auth — email + password, single-user app).")])
add_bullet([b("Onboarding wizard"), t(" (/onboarding, 6 steps + review):")])
add_bullet([b("Health"), t(" — age, sex, height, conditions, medications (GLP-1 status flagged here), injuries, equipment access.")], level=1)
add_bullet([b("Training"), t(" — experience, current frequency, recent e1RMs for Big-4 lifts.")], level=1)
add_bullet([b("Lifestyle"), t(" — job, stress, family obligations, travel cadence.")], level=1)
add_bullet([b("Nutrition"), t(" — current intake, dietary style, allergies, kitchen reality.")], level=1)
add_bullet([b("Sleep"), t(" — typical bed/wake, quality, environment.")], level=1)
add_bullet([b("Goals"), t(" — primary goal + the “why” narrative (durable client file).")], level=1)
add_bullet([b("Review & acknowledge"), t(" — locks the profile as v1 (immutable; future edits create v2).")], level=1)
add_bullet([b("Connect data sources"), t(" on /profile:")])
add_bullet([b("WHOOP"), t(" — OAuth → recovery, sleep, strain, HRV, RHR, SpO₂.")], level=1)
add_bullet([b("Withings"), t(" — OAuth → weight, body fat, muscle mass, hydration.")], level=1)
add_bullet([b("Apple Health (Garmin)"), t(" — generates a per-user bearer token; user installs the Shortcut that POSTs steps/calories/distance daily.")], level=1)
add_bullet([b("Strong / Yazio"), t(" — same bearer-token webhook (lifts + nutrition).")], level=1)
add_bullet([b("Generate plan"), t(" (CTA on /profile) → 5-beat chat intake on /coach → deterministic plan composer produces strength template, nutrition targets, sleep window, recovery protocol. User commits via approval token.")])
add_bullet([b("Set up training block"), t(" in /coach (mode = setup_block) — 5-week mesocycle goal (e.g. “+5 kg on squat 1RM, hold 80 kg LBM”).")])

# --- 2. Typical day
add_h2("2. A typical day — morning → evening")
add_bullet([b("Wake up"), t(" → open app → "), b("Morning intake"), t(" auto-launches as a chat:")])
add_bullet([t("Bot asks (structured chips, not free text): sick? fatigue? bloating? soreness areas + severity? sleep feel?")], level=1)
add_bullet([t("State machine writes answers to checkins.intake_state, then transitions to assembling_brief.")], level=1)
add_bullet([b("Morning brief card"), t(" appears in chat (single structured card, 5–7 blocks):")])
add_bullet([b("Yesterday recap"), t(" — sleep, strain, steps, macros hit/missed.")], level=1)
add_bullet([b("Readiness band"), t(" (high / moderate / low) with the top + and − drivers.")], level=1)
add_bullet([b("Today’s session"), t(" (training day: lifts + sets/reps + RPE target) "), b("or recovery focus"), t(" (rest day).")], level=1)
add_bullet([b("Macros target"), t(" (protein floor, calories, carbs) — sourced from active plan; GLP-1 mode bumps protein to 1.8–2.0 g/kg.")], level=1)
add_bullet([b("Coach advice"), t(" — one short AI-generated paragraph adapting to today’s flags (low sleep, alcohol last night, GLP-1 phase, missed protein streak…).")], level=1)
add_bullet([b("Tonight’s sleep target"), t(" (lights-out time).")], level=1)
add_bullet([b("Mid-morning"), t(" → if readiness is low and today is hard, brief surfaces a "), b("swap chip"), t(" (“swap today’s heavy lower for tomorrow’s pull?”) — one tap reorders the week.")])
add_bullet([b("Daytime"), t(":")])
add_bullet([t("Steps + active calories stream in automatically (Garmin → Apple Health → webhook).")], level=1)
add_bullet([t("Eat → Yazio logs meal → webhook updates today’s macros bar.")], level=1)
add_bullet([t("Train → log sets in Strong → CSV/webhook syncs → /strength tab shows the session vs plan, PRs, and the muscle-map highlight.")], level=1)
add_bullet([b("Coach speaks first"), t(" — proactive chat nudges may appear without you opening anything: e.g. \"Squat e1RM has been flat for 3 weeks — switch to a deload?\", or \"Weight is +0.4 kg/wk vs your 0.2 target — tighten the deficit?\". Each card deep-links into the relevant trend section. Triggers are deterministic: plateau detection, off-pace weight vs goal, HRV below baseline.")])
add_bullet([b("/strength tab"), t(" during/after workout: today’s planned session, swap sheet if life got in the way, exercise trend cards, click an exercise → highlights worked muscles on the anatomy diagram.")])
add_bullet([b("/log tab"), t(" anytime: manual entry for anything the integrations miss (mood, hydration, notes, weight if not on a Withings scale that day).")])
add_bullet([b("/coach tab"), t(": open chat for ad-hoc questions (“can I drop calories 100 more this week?”, “knee feels off, modify squat day?”) — coach answers with the full snapshot in context.")])
add_bullet([b("Evening"), t(":")])
add_bullet([t("Optional check-in (energy, mood, alcohol).")], level=1)
add_bullet([t("WHOOP auto-captures sleep overnight — feeds tomorrow’s recovery score (cron /api/whoop/sync runs 08:00 UTC).")], level=1)

# --- 3. Weekly + monthly
add_h2("3. Weekly + monthly cadence")
add_bullet([b("Sunday — Weekly Review document"), t(" (auto-generated each Sunday cron):")])
add_bullet([b("Recap"), t(" — what was committed Sunday vs what actually happened (as-planned / swapped / missed / rest), with prose distinguishing each status.")], level=1)
add_bullet([b("Reconfirm"), t(" — the coach asks 1–3 short questions surfacing tensions worth a human call (e.g. \"sleep efficiency dropped — is bedtime drifting?\").")], level=1)
add_bullet([b("Trends"), t(" — strength-per-LBM, allometric and IPF GL deltas, body-comp slope, plateau spans, cross-metric insights — each cell deep-links to the matching /coach/trends section.")], level=1)
add_bullet([b("Prescription + targets"), t(" — next week’s session plan, macros, sleep window, and recovery focus, ready to commit.")], level=1)
add_bullet([b("Three surfaces, one document"), t(": chat card on Sunday, full review page at /coach/weeks/[week_start], and a Tue-Sat banner on /coach if you haven’t answered yet. Versioned per (user, week_start) — explicit regenerate increments the version, prior draft superseded.")], level=1)
add_bullet([b("Commit"), t(" via approval token → writes the next training_weeks row and stamps the review with committed_training_week_id; strength tab anchors to that Sunday commitment all week.")], level=1)
add_bullet([b("Mid-week swaps"), t(" — A↔B exchange or single-day replacement; original Sunday plan preserved (original_session_plan jsonb) so adherence math stays anchored.")])
add_bullet([b("Monthly — Body measurements"), t(" (/log → measurements form): 14 circumferences + photo (private storage), tracks composition drift the scale can’t see.")])
add_bullet([b("Every 5 weeks"), t(": new training block — repeat block-setup ritual.")])

# --- 4. Coach analysis & proactive layer
add_h2("4. Coach analysis & proactive layer")
add_bullet([b("/coach/trends — the deep coaching surface"), t(" (distinct from raw /trends):")])
add_bullet([b("Performance"), t(" — per-lift e1RM trends with OLS slopes, plateau spans, strength-per-LBM and allometric scaling, IPF GL.")], level=1)
add_bullet([b("Composition"), t(" — weight, lean mass, fat mass slopes vs the active block’s targets (off-pace flags fire here).")], level=1)
add_bullet([b("Cross"), t(" — plain-English insights linking metrics (e.g. \"protein gap correlates with the plateau on incline press\"). Pure deterministic templating, no AI calls.")], level=1)
add_bullet([t("URL-driven section state (?section=performance|composition|cross) so chat cards and weekly-review cells deep-link straight to the relevant view.")], level=1)
add_bullet([b("Proactive nudges"), t(" — a daily 11:00 UTC cron evaluates 3 deterministic triggers against the same trend data and writes one chat card per fired trigger:")])
add_bullet([b("Plateau"), t(" — a Big-4 lift’s e1RM has been flat past the configured window.")], level=1)
add_bullet([b("Off-pace weight"), t(" — body-weight slope is materially off the block’s target (cut/bulk/recomp).")], level=1)
add_bullet([b("HRV below baseline"), t(" — rolling HRV has dropped vs the WHOOP baseline.")], level=1)
add_bullet([t("Cards are pure templating, deduped via the chat history within a 7-day window — the chat thread itself is the audit trail. No push notifications (yet) — they live in /coach so the next time you open chat they’re waiting.")], level=1)
add_bullet([b("/ dashboard"), t(" — today’s readiness, top drivers, quick-glance week strip.")])
add_bullet([b("/trends"), t(" — raw metric explorer: any field over 1W / 1M / LY / YTD / custom range, filter by source.")])

# --- Mental model footer
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("The mental model")
run.bold = True
run.font.size = Pt(10)
run2 = p.add_run(
    ": the app is a coach that already knows your file (acknowledged profile + active plan), watches "
    "the streams (WHOOP, Withings, Apple Health, Strong, Yazio), and intercepts you on three rhythms — "
    "morning brief sets the day, the proactive layer reaches out when a trigger fires, and the Sunday "
    "review document closes the week. Everything else is the system removing friction from data entry "
    "while keeping the human decisions (commit week, commit block, acknowledge milestones) in your hands."
)
run2.font.size = Pt(10)

doc.save(OUT)
print(f"Wrote {OUT}")
